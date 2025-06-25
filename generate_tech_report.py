import os
from pathlib import Path

from docx import Document
from docx.shared import Inches
try:
    from graphviz import Digraph
    _GRAPHVIZ_OK = True
except ImportError:  # pragma: no cover
    Digraph = None  # type: ignore
    _GRAPHVIZ_OK = False


# ---------------------------------------------------------------------------
# Helper: ensure output / assets directories
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).parent
OUT_DIR = BASE_DIR / "reports"
IMG_DIR = OUT_DIR / "images"
OUT_DIR.mkdir(exist_ok=True)
IMG_DIR.mkdir(exist_ok=True)


# ---------------------------------------------------------------------------
# 1. Build pipeline diagram (Graphviz)
# ---------------------------------------------------------------------------

def _build_pipeline_diagram(out_path: Path) -> None:
    """Create a simple left-to-right Graphviz diagram of the 5-stage QA pipeline."""
    if not _GRAPHVIZ_OK:
        return
    g = Digraph(name="pipeline", comment="Quran QA Pipeline", format="png")
    g.attr(rankdir="LR", fontsize="11", fontname="Helvetica")

    g.node("Q", "User Question", shape="oval", style="filled", fillcolor="#E8F0FE")
    g.node("C", "Stage 1:\nClassification", shape="box")
    g.node("E", "Stage 2:\nEntity Extraction", shape="box")
    g.node("R", "Stage 3:\nContext Retrieval", shape="box")
    g.node("P", "Stage 4:\nPrompt Builder", shape="box")
    g.node("L", "Stage 5:\nLLM Answer", shape="box")
    g.node("A", "Answer", shape="oval", style="filled", fillcolor="#E8F0FE")

    g.edges([
        ("Q", "C"),
        ("C", "E"),
        ("E", "R"),
        ("R", "P"),
        ("P", "L"),
        ("L", "A"),
    ])

    g.render(out_path, cleanup=True)


def _build_extractor_diagram(out_path: Path) -> None:
    """Create a decision-tree diagram for the extractor's layered logic."""
    if not _GRAPHVIZ_OK:
        return
    g = Digraph(name="extractor", comment="Entity-Extraction Flow", format="png")
    g.attr(rankdir="TB", fontsize="11", fontname="Helvetica")

    g.node("start", "Question Text", shape="oval", style="filled", fillcolor="#E8F0FE")
    g.node("regex", "Regex Layer", shape="box")
    g.node("regex_ok", "Word(s) found", shape="oval", style="filled", fillcolor="#C6F6D5")
    g.node("llm", "LLM Fallback", shape="box")
    g.node("llm_ok", "Word(s) found", shape="oval", style="filled", fillcolor="#C6F6D5")
    g.node("fail", "Ask user for clarification", shape="octagon", style="filled", fillcolor="#FDE68A")

    g.edge("start", "regex")
    g.edge("regex", "regex_ok", label="match", fontsize="10")
    g.edge("regex", "llm", label="no match", fontsize="10")
    g.edge("llm", "llm_ok", label=">= 0.45 conf", fontsize="10")
    g.edge("llm", "fail", label="< 0.45 conf", fontsize="10")

    g.render(out_path, cleanup=True)


# ---------------------------------------------------------------------------
# 2. Assemble DOCX Report
# ---------------------------------------------------------------------------

def build_report(doc_path: Path) -> None:
    doc = Document()

    # Title page
    doc.add_heading("Quran Chatbot – Comprehensive Technical Report", 0)
    doc.add_paragraph("Version: 1.0")
    doc.add_paragraph("Generated automatically by generate_tech_report.py\n")

    # 1. Overview
    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "The project provides an interactive Arabic Q&A assistant specialised in Qur'anic linguistics. "
        "It uses a multi-stage pipeline that classifies the user's question, extracts target linguistic "
        "entities (word / root / sūrah), gathers just-enough contextual data from local JSONL corpora, "
        "and finally crafts a rich prompt for a large-language model (LLM) hosted via OpenAI's API."
    )

    # 2. High-level architecture
    doc.add_heading("2. High-level Architecture", level=1)
    pipeline_png = IMG_DIR / "pipeline"
    if _GRAPHVIZ_OK:
        _build_pipeline_diagram(pipeline_png)
        doc.add_picture(str(pipeline_png.with_suffix(".png")), width=Inches(5.5))
        doc.add_paragraph(
            "Figure 1 – End-to-end dataflow. Each stage is implemented in its own module and invoked by "
            "pipeline.QuranQAPipeline.answer_question()."
        )
    else:
        doc.add_paragraph("[Graphviz not installed – diagram omitted]", style="Intense Quote")

    # 3. Detailed Stage Descriptions
    doc.add_heading("3. Pipeline Stages", level=1)

    # Stage 1
    doc.add_heading("3.1 Stage-1 – Question Classification", level=2)
    doc.add_paragraph(
        "Module: services/classification.py. Utilises a single chat-completion call to an LLM (default: gpt-4o-mini). "
        "The system prompt enumerates 12 canonical classes such as meaning_word, frequency_word_root, difference_two_words, etc. "
        "The assistant must reply with the bare slug which is validated against a whitelist."
    )

    # Stage 2
    doc.add_heading("3.2 Stage-2 – Target Entity Extraction", level=2)
    doc.add_paragraph(
        "Module: services/extractors/quranic_word_extractor.py + surah_extractor.py. A layered approach is employed:\n"
        "1. **Regex heuristics** (see Table 1) search for explicit mentions of Qur'anic words, roots or two-word 'difference' questions.\n"
        "2. **LLM fallback** returns a JSON payload {word, confidence}. A threshold of 0.45 decides acceptance.\n"
        "3. Failure to identify an entity triggers an explicit clarification message."
    )
    extract_png = IMG_DIR / "extractor"
    if _GRAPHVIZ_OK:
        _build_extractor_diagram(extract_png)
        doc.add_picture(str(extract_png.with_suffix(".png")), width=Inches(4.5))
    else:
        doc.add_paragraph("[Graphviz not installed – extractor diagram omitted]", style="Intense Quote")

    # Regex table
    doc.add_paragraph("Table 1 – Principal regex patterns for single-word extraction (simplified):", style="Intense Quote")
    rex_list = [
        r"(?:(?:ما\s+)?(?:معنى|تفسير|مدلول)\s+(?:من\s+|ب)?(?:كلمة|لفظة)\s+)([^\s\?\.،؟]+)",
        r"(?:(?:ماذا\s+)?يعني\s+(?:لفظة?|كلمة)\s+)([^\s\?\.،؟]+)",
        r"(?:(?:فسر|اشرح)\s+(?:ال)?(?:كلمة|لفظة)\s+)([^\s\?\.،؟]+)",
        r"(?:جذر|اشتقاق)\s+(?:ال)?(?:كلمة|لفظة)\s+([^\s\?\.،؟]+)",
        r"(?:تصريفات)?\s*(?:جذر)\s+([^\s\?\.،؟]+)",
        r"(?:كلمة|لفظة)\s+([أإآ][^\s\?\.،؟]+)",
        r"(?:كلمة|لفظة)\s+([^\s\?\.،؟]+)",
    ]
    for pat in rex_list:
        doc.add_paragraph(pat, style="List Bullet")

    # Stage 3
    doc.add_heading("3.3 Stage-3 – Context Retrieval", level=2)
    doc.add_paragraph(
        "Module: services/retrievers/dispatcher.py. Implements a 12-way branching strategy that calls "
        "specialised helpers (morphology_retriever, root_retriever, dictionary_retriever, etc.). It leverages a fast on-disk "
        "JSONL morphology corpus (~100 KB gzipped) and uses the smart_exact_match() utility to tolerate proclitics, definite articles "
        "and hamza variants."
    )

    # Stage 4
    doc.add_heading("3.4 Stage-4 – Prompt Building", level=2)
    doc.add_paragraph(
        "Module: pipeline/prompt_builder.py crafts a role-annotated chat prompt (<question>, <context>) plus extra instructions "
        "depending on the detected question-type. A UTC timestamp is appended for traceability."
    )

    # Stage 5
    doc.add_heading("3.5 Stage-5 – LLM Query & Answer Synthesis", level=2)
    doc.add_paragraph(
        "Module: services/llm.py. Sends the prompt to the configured ChatCompletion model and streams verbose logs "
        "if the UI requests it. The answer is returned verbatim to the UI (CLI or Streamlit)."
    )

    # 4. Data Resources
    doc.add_heading("4. Data Resources", level=1)
    doc.add_paragraph("The system relies on three lightweight JSONL files located under data/:")
    doc.add_paragraph("quran_morphology.jsonl – 77 K token records comprising lemma, root and positional metadata.", style="List Bullet")
    doc.add_paragraph("root_analysis.jsonl – Compact list of ≈ 2 K triliteral roots with semantic notes.", style="List Bullet")
    doc.add_paragraph("arabic_dictionary.jsonl – General-purpose Arabic dictionary dump used for fallback definitions.", style="List Bullet")

    # 5. Utilities & Key Algorithms
    doc.add_heading("5. Utilities & Key Algorithms", level=1)
    doc.add_paragraph("utils/arabic.normalize() strips diacritics, harmonises hamza/alif variants and preserves shadda.")
    doc.add_paragraph("smart_exact_match() (services/retrievers/morphology_retriever.py) yields robust lemma / surface matches by generating orthographic variants.")
    doc.add_paragraph("embedding_utils.get_embeddings() provides SBERT embeddings ready for vector search tasks (reserved for future semantic retrievers).")

    # 6. Interfaces
    doc.add_heading("6. Interfaces", level=1)
    doc.add_paragraph("CLI: python main.py \"…\"")
    doc.add_paragraph("Web UI: Streamlit app.py – real-time status callback renders the pipeline trace.")

    # 7. Security & Rate-limiting
    doc.add_heading("7. Operational Considerations", level=1)
    doc.add_paragraph("• OpenAI API key is loaded via dotenv; never hard-coded.\n• Token usage is capped (max_tokens).\n• Morphology files are read-only, eliminating risk of injection.")

    # 8. Roadmap
    doc.add_heading("8. Future Work", level=1)
    doc.add_paragraph("– Add caching layer for root_ayah_extraction.\n– Integrate local open-source LLMs.\n– Expand regex library for additional Arabic patterns.")

    # Save file
    doc.save(str(doc_path))
    print(f"✅ Report generated at {doc_path}")


if __name__ == "__main__":
    output_doc = OUT_DIR / "Quran_Chatbot_Technical_Report_Auto.docx"
    build_report(output_doc) 